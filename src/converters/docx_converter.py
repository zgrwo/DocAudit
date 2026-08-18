"""DOCX 转换器 — 使用 python-docx 解析，保留段落/字符样式信息"""

import logging
import re
from collections import defaultdict
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Length

from src.converters.base import BaseConverter
from src.models.document import (
    Document,
    DocumentMetadata,
    Page,
    PageElement,
    Paragraph,
    Run,
    TableCell,
)

logger = logging.getLogger(__name__)

ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

# 样式型标题回退 (HIGH-1): "Heading N" 样式名 → 0-based 层级 N-1
HEADING_STYLE_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


class DocxConverter(BaseConverter):
    """将 DOCX 文件转换为统一 Document 模型。

    保留信息：
    - 段落样式名称 + 大纲级别
    - Run 级字体/字号/加粗/斜体/颜色
    - 段落对齐方式 + 间距
    - 表格结构
    - 文档元素顺序（段落和表格交错出现）
    """

    def can_handle(self, source_path: str | Path) -> bool:
        ext = Path(source_path).suffix.lower().lstrip(".")
        return ext in ("docx", "doc")

    def convert(self, source_path: str | Path) -> Document:
        source_path = Path(source_path)
        logger.info("解析 DOCX: %s", source_path.name)

        try:
            doc = DocxDocument(str(source_path))
        except Exception as e:
            logger.error("无法打开 DOCX 文件: %s", e)
            raise ValueError(f"无法解析 DOCX 文件: {source_path}") from e

        # --- 元数据 ---
        metadata = self._extract_metadata(doc)

        # --- 构建元素查找表 (局部变量, 线程安全) ---
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        # --- 遍历文档主体（保留段落和表格的交错顺序） ---
        elements: list[PageElement] = []

        for child in doc.element.body:
            tag = _tag_name(child)

            if tag == "p":
                para = para_map.get(child)
                if para is not None:
                    elem = self._convert_paragraph(para)
                    if elem is not None:
                        elements.append(elem)

            elif tag == "tbl":
                tbl = table_map.get(child)
                if tbl is not None:
                    elem = self._convert_table(tbl)
                    if elem is not None:
                        elements.append(elem)

            elif tag == "sdt":
                # SDT 结构: sdtPr → sdtContent → sdtEndPr
                # 只进入 sdtContent (sdtPr 仅含属性，不含内容段落)
                # 注意: 嵌套 SDT (SDT 内含 SDT) 当前仅处理一层，罕见场景见 agents.md 已知陷阱 #6
                for inner in child:
                    inner_tag = _tag_name(inner)
                    if inner_tag == "sdtContent":
                        for item in inner:
                            item_tag = _tag_name(item)
                            if item_tag == "p":
                                para = para_map.get(item)
                                if para is not None:
                                    elem = self._convert_paragraph(para)
                                    if elem is not None:
                                        elements.append(elem)
                            elif item_tag == "tbl":
                                tbl = table_map.get(item)
                                if tbl is not None:
                                    elem = self._convert_table(tbl)
                                    if elem is not None:
                                        elements.append(elem)

        # 未解析嵌套内容提示 (仅日志，不改变解析行为)
        self._warn_unparsed_nested_content(doc)

        logger.info("DOCX 解析完毕: %d 元素", len(elements))

        # 分页
        pages = self._split_into_pages(elements)

        return Document(
            source_path=str(source_path),
            format="docx",
            metadata=metadata,
            pages=pages,
        )

    def _warn_unparsed_nested_content(self, doc: DocxDocument) -> None:
        """提示当前未解析的嵌套内容：文本框、页眉页脚、脚注。

        仅输出 logger.warning 说明跳过范围，不改变现有解析行为。
        """
        # 文本框: 段落内 w:drawing/w:txbxContent (画布文本框中的文字)
        try:
            for para in doc.paragraphs:
                if para._element.findall(".//" + qn("w:txbxContent")):
                    logger.warning(
                        "段落包含文本框 (w:txbxContent) 内容，当前跳过未解析: %s",
                        para.text[:40] or "(无文本)",
                    )
        except Exception:
            pass  # bare-handler-ok — 文本框检测降级，不影响解析

        # 页眉/页脚
        try:
            for section in doc.sections:
                for hf in (section.header, section.footer):
                    if not hf.is_linked_to_previous and any(p.text.strip() for p in hf.paragraphs):
                        logger.warning("检测到页眉/页脚内容，当前跳过未解析（不在正文审查范围内）")
                        break  # 每节最多一条
        except Exception:
            pass  # bare-handler-ok — 页眉页脚检测降级，不影响解析

        # 脚注: 文档部件存在 footnotes 关系即视为含脚注
        try:
            has_footnotes = any("footnotes" in rel.reltype for rel in doc.part.rels.values())
            if has_footnotes:
                logger.warning("检测到脚注内容，当前跳过未解析")
        except Exception:
            pass  # bare-handler-ok — 脚注检测降级，不影响解析

    def _extract_metadata(self, doc: DocxDocument) -> DocumentMetadata:
        """提取文档元数据"""
        # 统计字数 (在 try/except 外计算，避免重复)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
        try:
            props = doc.core_properties
            return DocumentMetadata(
                title=props.title,
                author=props.author,
                created=str(props.created) if props.created else None,
                modified=str(props.modified) if props.modified else None,
                word_count=word_count,
            )
        except (AttributeError, ValueError) as e:
            logger.debug("元数据提取失败: %s", e, exc_info=True)
            return DocumentMetadata(word_count=word_count)

    def _convert_paragraph(self, para) -> PageElement | None:
        """转换单个段落"""
        runs: list[Run] = []
        for r in para.runs:
            try:
                font = r.font
                font_color = None
                try:
                    if font.color and font.color.rgb:
                        font_color = str(font.color.rgb)
                except (AttributeError, ValueError):
                    pass

                # eastAsia 中文字体: w:rPr/w:rFonts 的 w:eastAsia 属性
                # (python-docx font.name 只读 w:ascii/w:hAnsi，不含中文显示字体)
                font_name_east_asia = None
                try:
                    rPr = r._element.rPr
                    if rPr is not None and rPr.rFonts is not None:
                        font_name_east_asia = rPr.rFonts.get(qn("w:eastAsia"))
                except Exception:
                    pass  # bare-handler-ok — eastAsia 字体提取降级，失败时保留 None

                runs.append(
                    Run(
                        text=r.text,
                        font_name=font.name,
                        font_name_east_asia=font_name_east_asia,
                        font_size=font.size.pt if font.size else None,
                        bold=font.bold,
                        italic=font.italic,
                        underline=font.underline,
                        color=font_color,
                    )
                )
            except Exception as e:
                # F9: 单个 run 提取失败仅跳过该 run (不丢整段格式)，并继续处理其余 run
                logger.debug("DOCX run 格式提取失败，跳过该 run: %s", e, exc_info=True)
                continue

        # 大纲级别
        level = None
        try:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                outline = pPr.find(qn("w:outlineLvl"))
                if outline is not None:
                    level = int(outline.get(qn("w:val"), "0"))
        except Exception:
            pass  # bare-handler-ok — 大纲级别提取降级，失败时保留 None

        full_text = para.text
        if not full_text.strip() and not runs:
            return None

        # 行距: 固定行距返回 Twips (Length, int 子类)，转为 pt 以满足 float|None 契约 (F5)
        line_spacing = None
        if para.paragraph_format and para.paragraph_format.line_spacing is not None:
            line_spacing = para.paragraph_format.line_spacing
            if isinstance(line_spacing, Length):
                line_spacing = line_spacing.pt

        # 判断是否为标题（Heading 样式）
        is_heading = False
        style_name = None
        try:
            style_name = para.style.name
            is_heading = style_name and "heading" in style_name.lower()
        except Exception:
            pass  # bare-handler-ok — 样式名读取降级，失败时按非标题处理

        # 样式级回退 (HIGH-1): 段落级无 outlineLvl 时，从样式定义补 level。
        # ① style 的 w:pPr/w:outlineLvl (python-docx add_heading 生成的样式型标题即此形态)
        # ② 样式名 "Heading N" → N-1 (样式未定义 outlineLvl 时)
        if is_heading and level is None:
            try:
                style_el = para.style._element
                if style_el is not None:
                    style_pPr = style_el.find(qn("w:pPr"))
                    if style_pPr is not None:
                        outline = style_pPr.find(qn("w:outlineLvl"))
                        if outline is not None:
                            level = int(outline.get(qn("w:val"), "0"))
            except Exception:
                pass  # bare-handler-ok — 样式级 outlineLvl 提取降级，失败时保留 None
            if level is None and style_name:
                m = HEADING_STYLE_RE.match(style_name)
                if m:
                    level = int(m.group(1)) - 1

        return PageElement(
            type="text_frame",
            paragraphs=[
                Paragraph(
                    text=full_text,
                    runs=runs,
                    level=level,
                    alignment=ALIGNMENT_MAP.get(para.alignment),
                    space_before=(
                        para.paragraph_format.space_before.pt
                        if (
                            para.paragraph_format and para.paragraph_format.space_before is not None
                        )
                        else None
                    ),
                    space_after=(
                        para.paragraph_format.space_after.pt
                        if (para.paragraph_format and para.paragraph_format.space_after is not None)
                        else None
                    ),
                    line_spacing=line_spacing,
                )
            ],
            # 段落样式名放入 style_name，shape_name 保留给 PPTX shape 名 (默认 None)
            style_name=style_name,
            # H1 或 H2 均视为标题 (与 _split_into_pages 的页面边界语义一致: level <= 1)
            is_title=is_heading and level is not None and level <= 1,
            is_body=not is_heading,
        )

    def _convert_table(self, table) -> PageElement | None:
        """转换表格"""
        if not table.rows:
            return None  # 空表格 — 无有意义内容
        try:
            cells: list[TableCell] = []
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    # 提取第一个段落第一个 run 的字体信息
                    font_name = None
                    font_size = None
                    font_color = None
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        first_run = cell.paragraphs[0].runs[0]
                        font_name = first_run.font.name
                        if first_run.font.size:
                            font_size = first_run.font.size.pt
                        try:
                            if first_run.font.color and first_run.font.color.rgb:
                                font_color = str(first_run.font.color.rgb)
                        except (AttributeError, ValueError):
                            font_color = None  # 主题色等无法取 RGB — 降级

                    # 单元格底色 (w:shd w:fill 直接指定的纯色；继承样式/无底色 → None)
                    fill_color = None
                    try:
                        tcPr = cell._tc.tcPr
                        if tcPr is not None:
                            shd = tcPr.find(qn("w:shd"))
                            if shd is not None:
                                fill = shd.get(qn("w:fill"))
                                if fill and fill.lower() != "auto":
                                    fill_color = fill.upper()
                    except Exception:
                        fill_color = None  # bare-handler-ok — 异常 XML 结构，降级

                    cells.append(
                        TableCell(
                            text=cell_text,
                            row=row_idx,
                            col=col_idx,
                            font_name=font_name,
                            font_size=font_size,
                            fill_color=fill_color,
                            font_color=font_color,
                        )
                    )

            nrows = len(table.rows)
            # 按行分组 (单次遍历，与 PptxConverter._convert_table 一致)
            row_map: dict[int, list[TableCell]] = defaultdict(list)
            for c in cells:
                row_map[c.row].append(c)
            rows: list[list[TableCell]] = [row_map[r] for r in sorted(row_map) if r < nrows]

            return PageElement(
                type="table",
                tables=rows,
            )
        except Exception as e:
            logger.warning("表格转换失败: %s", e)
            return None

    def _split_into_pages(self, elements: list[PageElement]) -> list[Page]:
        """按语义边界拆分为多页：优先以标题(H1/H2)为页面边界，CHUNK_SIZE 为硬上限回退。

        这避免了将标题切到前一页末尾或后一页开头导致的假跳级问题。
        """
        pages: list[Page] = []
        CHUNK_SIZE = 40  # 硬上限回退 — 仅在无标题可用作页面边界时触发

        # 空文档直接返回单页空内容
        if not elements:
            return [Page(index=0, elements=[], slide_number=1)]

        current_chunk: list[PageElement] = []
        chunk_split_count = 0  # 统计 CHUNK_SIZE 回退触发次数，便于调试

        def _is_heading(elem: PageElement) -> bool:
            """判断元素是否为 H1 或 H2 级别标题"""
            if elem.type != "text_frame" or not elem.paragraphs:
                return False
            level = elem.paragraphs[0].level
            return level is not None and level <= 1

        for elem in elements:
            # 标题作为页面边界（但不在 chunk 为空时创建空页）
            if _is_heading(elem) and current_chunk:
                pages.append(
                    Page(
                        index=len(pages),
                        elements=current_chunk,
                        slide_number=len(pages) + 1,
                    )
                )
                current_chunk = []
            current_chunk.append(elem)
            # CHUNK_SIZE 硬上限回退：避免单页元素过多
            if len(current_chunk) >= CHUNK_SIZE:
                chunk_split_count += 1
                logger.debug(
                    "DOCX 分页: CHUNK_SIZE (%d) 回退触发 (第 %d 次)，当前页 %d 个元素",
                    CHUNK_SIZE,
                    chunk_split_count,
                    len(current_chunk),
                )
                pages.append(
                    Page(
                        index=len(pages),
                        elements=current_chunk,
                        slide_number=len(pages) + 1,
                    )
                )
                current_chunk = []

        if current_chunk:
            pages.append(
                Page(
                    index=len(pages),
                    elements=current_chunk,
                    slide_number=len(pages) + 1,
                )
            )

        return pages


def _tag_name(element) -> str:
    """获取 XML 元素的标签名（去除命名空间）"""
    tag = element.tag
    return tag.split("}")[-1] if "}" in tag else tag
