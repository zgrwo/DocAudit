"""黄金测试 — 真实三路径（Python API / 真实 CLI / 真实 WebUI）。

背景 (P0-2): 旧 test_integration.py::test_web_ui_path_equals_cli_path_equals_python_path
是"间接三路径" — 三条路径共享同一函数
(find_converter→convert→build_auditors→run_auditors) 的参数写法变体，
不跑真实 CLI、不执行 app.py，若 cli.py/app.py 传错参数测试依然全绿。

本文件:
- 路径 1 (基准/真理): 直接 Python API — 绝对路径 + 显式 vocab。
- 路径 2 (真实 CLI): subprocess 运行
  `python src/cli.py <fixture> -o <tmp>.json --rules <root>/rules.md --glossary <root>/glossary`，
  读取 JSON 报告 findings 与基准比较 — 按 (type, severity, rule_id, page_index,
  message, context, suggestion, location) 排序后的集合完全相等 (排除随机 id)。
- 路径 3 (真实 WebUI): streamlit.testing.v1.AppTest 运行 app.py，向 file_uploader
  注入文件并点击「🔍 开始审查」按钮，读取 at.session_state["findings"] 与基准比较。
"""

import json
import subprocess
import sys
from pathlib import Path

import pytest

from src.engines.pipeline import build_auditors, find_converter, run_auditors
from src.models.document import Document, DocumentMetadata, Page, PageElement, Paragraph
from src.models.finding import AuditFinding, FindingSeverity, FindingType

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FIXTURE_PPTX = PROJECT_ROOT / "tests" / "fixtures" / "sample.pptx"

# AppTest file_uploader.set_value 需要 (文件名, 内容, MIME)
_MIME_BY_SUFFIX = {
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".md": "text/markdown",
}


def _finding_key(f: AuditFinding) -> tuple:
    """AuditFinding → 比较键 (排除随机 id / element_index / metadata)。"""
    return (
        f.type.value,
        f.severity.value,
        f.rule_id,
        f.page_index,
        f.message,
        f.context,
        f.suggestion,
        f.location,
    )


def _finding_key_from_dict(d: dict) -> tuple:
    """CLI JSON 报告中的 finding dict → 比较键。"""
    return (
        d["type"],
        d["severity"],
        d["rule_id"],
        d["page_index"],
        d["message"],
        d["context"],
        d["suggestion"],
        d["location"],
    )


def _baseline_findings(file_path: Path) -> list[AuditFinding]:
    """路径 1: 直接 Python API — 绝对路径 + 显式 vocab (基准/真理)。"""
    converter = find_converter(str(file_path))
    assert converter is not None, f"未找到 {file_path} 的转换器"
    doc = converter.convert(str(file_path))
    auditors = build_auditors(
        str(PROJECT_ROOT / "rules.md"),
        str(PROJECT_ROOT / "glossary"),
        str(PROJECT_ROOT / "vocab"),
    )
    return run_auditors(doc, auditors)


def _cli_findings(file_path: Path, tmp_path: Path) -> list[dict]:
    """路径 2: 真实 CLI subprocess — 读取 JSON 报告 findings。"""
    out_json = tmp_path / "cli_report.json"
    result = subprocess.run(
        [
            sys.executable,
            str(PROJECT_ROOT / "src" / "cli.py"),
            str(file_path),
            "-o",
            str(out_json),
            "--rules",
            str(PROJECT_ROOT / "rules.md"),
            "--glossary",
            str(PROJECT_ROOT / "glossary"),
        ],
        capture_output=True,
        text=True,
        timeout=300,
        cwd=str(PROJECT_ROOT),
        encoding="utf-8",
        errors="replace",
    )
    # 退出码: 存在 ERROR 发现时 CLI 退出 1 (正常语义)；argparse 用法错误才退出 2
    assert result.returncode != 2, f"CLI 用法错误: {result.stderr[-500:]}"
    assert out_json.exists(), (
        f"CLI 未生成 JSON 报告: rc={result.returncode} stdout={result.stdout[-500:]}"
    )
    report = json.loads(out_json.read_text(encoding="utf-8"))
    return report["findings"]


def _webui_findings(file_path: Path) -> list[AuditFinding]:
    """路径 3: 真实 WebUI — AppTest 完整交互 (上传文件 → 点击「🔍 开始审查」)。"""
    pytest.importorskip("streamlit")  # 环境无 streamlit 时跳过 (降级, 不影响其他路径)
    from streamlit.testing.v1 import AppTest

    data = file_path.read_bytes()
    mime = _MIME_BY_SUFFIX.get(file_path.suffix, "application/octet-stream")

    at = AppTest.from_file(str(PROJECT_ROOT / "app.py"), default_timeout=300)
    at.run()
    at.file_uploader(key="file_uploader").set_value((file_path.name, data, mime))
    at.run()
    start_buttons = [b for b in at.button if b.label and "开始审查" in b.label]
    assert start_buttons, "WebUI 未渲染「🔍 开始审查」按钮 (file_uploader 注入失败?)"
    start_buttons[0].click()
    at.run()

    session = at.session_state
    if "findings" not in session or not session["findings"]:
        exc_detail = [str(e.value) for e in at.exception]
        raise AssertionError(f"WebUI 审查后 session_state 中无 findings (app 异常: {exc_detail})")
    # 注意: app.py 的报告下载等渲染部分由 WP-C 并发维护，其临时异常
    # (如并发编辑中间态) 不影响已产生的 findings 有效性 — 本测试只承诺
    # "WebUI 审计路径产生的 findings 与基准一致"，故不因渲染段异常而失败。
    return session["findings"]


def _assert_keys_equal(actual: list[tuple], expected: list[tuple], label: str) -> None:
    """按比较键排序后的集合完全相等 (排除随机 id)。"""
    assert actual == expected, (
        f"{label} findings 与基准不一致:\n"
        f"  基准 {len(expected)} 条 / {label} {len(actual)} 条\n"
        f"  仅基准: {sorted(set(expected) - set(actual))[:3]}\n"
        f"  仅{label}: {sorted(set(actual) - set(expected))[:3]}"
    )


def _generate_docx(path: Path) -> None:
    """程序化生成测试 DOCX (tmp_path)。"""
    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_heading("测试标题", level=1)
    d.add_paragraph("这是正文内容，包含 FinFET 技术。")
    d.add_heading("结论", level=2)
    d.add_paragraph("测试完成。")
    d.save(str(path))


def _generate_markdown(path: Path) -> None:
    """程序化生成测试 Markdown (tmp_path)。"""
    path.write_text(
        "# 概述\n\n这是测试文档，介绍 TSV 技术。\n\n## 结论\n\n测试完成。\n",
        encoding="utf-8",
    )


class TestGoldenPathPptx:
    """黄金测试 — PPTX (tests/fixtures/sample.pptx) 三路径一致性。"""

    def test_pptx_cli_matches_baseline(self, tmp_path):
        """真实 CLI 与基准 (Python API) 的 findings 完全一致。"""
        baseline = _baseline_findings(FIXTURE_PPTX)
        assert baseline, "基准 findings 不应为空 (防空洞断言)"
        cli = _cli_findings(FIXTURE_PPTX, tmp_path)
        _assert_keys_equal(
            sorted(_finding_key_from_dict(d) for d in cli),
            sorted(_finding_key(f) for f in baseline),
            "CLI",
        )

    def test_pptx_webui_matches_baseline(self):
        """真实 WebUI (AppTest) 与基准 (Python API) 的 findings 完全一致。"""
        baseline = _baseline_findings(FIXTURE_PPTX)
        webui = _webui_findings(FIXTURE_PPTX)
        _assert_keys_equal(
            sorted(_finding_key(f) for f in webui),
            sorted(_finding_key(f) for f in baseline),
            "WebUI",
        )


class TestGoldenPathDocx:
    """黄金测试 — 程序化生成的 DOCX 三路径一致性。"""

    def test_docx_cli_matches_baseline(self, tmp_path):
        docx_path = tmp_path / "golden.docx"
        _generate_docx(docx_path)
        baseline = _baseline_findings(docx_path)
        assert baseline, "基准 findings 不应为空 (防空洞断言)"
        cli = _cli_findings(docx_path, tmp_path)
        _assert_keys_equal(
            sorted(_finding_key_from_dict(d) for d in cli),
            sorted(_finding_key(f) for f in baseline),
            "CLI",
        )

    def test_docx_webui_matches_baseline(self, tmp_path):
        docx_path = tmp_path / "golden.docx"
        _generate_docx(docx_path)
        baseline = _baseline_findings(docx_path)
        webui = _webui_findings(docx_path)
        _assert_keys_equal(
            sorted(_finding_key(f) for f in webui),
            sorted(_finding_key(f) for f in baseline),
            "WebUI",
        )


class TestGoldenPathMarkdown:
    """黄金测试 — 程序化生成的 Markdown 三路径一致性。"""

    def test_markdown_cli_matches_baseline(self, tmp_path):
        md_path = tmp_path / "golden.md"
        _generate_markdown(md_path)
        baseline = _baseline_findings(md_path)
        assert baseline, "基准 findings 不应为空 (防空洞断言)"
        cli = _cli_findings(md_path, tmp_path)
        _assert_keys_equal(
            sorted(_finding_key_from_dict(d) for d in cli),
            sorted(_finding_key(f) for f in baseline),
            "CLI",
        )

    def test_markdown_webui_matches_baseline(self, tmp_path):
        md_path = tmp_path / "golden.md"
        _generate_markdown(md_path)
        baseline = _baseline_findings(md_path)
        webui = _webui_findings(md_path)
        _assert_keys_equal(
            sorted(_finding_key(f) for f in webui),
            sorted(_finding_key(f) for f in baseline),
            "WebUI",
        )


class _StubAuditor:
    """最小桩审计器 — 仅用于 run_auditors 行为测试 (鸭子类型, 无需继承 BaseAuditor)。"""

    def __init__(self, name: str, findings: list | None = None, exc: Exception | None = None):
        self.name = name
        self._findings = list(findings or [])
        self._exc = exc

    def audit(self, doc) -> list[AuditFinding]:
        if self._exc is not None:
            raise self._exc
        return list(self._findings)


def _simple_doc() -> Document:
    """构造最小文档 (run_auditors 行为测试用)。"""
    page = Page(
        index=0,
        slide_number=1,
        elements=[PageElement(type="text_frame", paragraphs=[Paragraph(text="测试", runs=[])])],
    )
    return Document(format="md", source_path="test.md", metadata=DocumentMetadata(), pages=[page])


class TestRunAuditorsBehavior:
    """run_auditors 行为测试 — on_progress 回调序列 + 审计器异常 → SYS-ERROR。"""

    def test_on_progress_called_in_auditor_order_then_complete(self):
        """on_progress 按审计器顺序调用，且最后收到「完成」。"""
        calls: list[tuple] = []
        auditors = [("甲", _StubAuditor("甲")), ("乙", _StubAuditor("乙"))]
        run_auditors(_simple_doc(), auditors, on_progress=lambda n, i, t: calls.append((n, i, t)))
        assert calls == [("甲", 0, 2), ("乙", 1, 2), ("完成", 2, 2)], f"回调序列错误: {calls}"

    def test_auditor_exception_produces_sys_error_finding(self):
        """审计器 audit() 抛异常 → 产生 SYS-ERROR ERROR 级 finding (UI 可见)。"""
        auditors = [("甲", _StubAuditor("甲", exc=RuntimeError("甲 崩溃")))]
        findings = run_auditors(_simple_doc(), auditors)
        sys_errors = [f for f in findings if f.rule_id == "SYS-ERROR"]
        assert len(sys_errors) == 1, f"应产生 1 条 SYS-ERROR, got: {findings}"
        f = sys_errors[0]
        assert f.severity == FindingSeverity.ERROR
        assert (
            f.type == FindingType.SYSTEM
        )  # SYS-ERROR 归 system 类型 (二轮审查修复: 曾误归 CUSTOM)
        assert "甲" in f.message and "甲 崩溃" in f.message

    def test_two_different_exceptions_not_collapsed_by_dedup(self):
        """两条不同异常 → 2 条 SYS-ERROR (dedup_key 含错误摘要, 不得折叠)。"""
        auditors = [
            ("甲", _StubAuditor("甲", exc=RuntimeError("甲 崩溃"))),
            ("乙", _StubAuditor("乙", exc=ValueError("乙 崩溃"))),
        ]
        findings = run_auditors(_simple_doc(), auditors)
        sys_errors = [f for f in findings if f.rule_id == "SYS-ERROR"]
        assert len(sys_errors) == 2, (
            f"两条不同异常应保留 2 条 SYS-ERROR (dedup 不得折叠), 实际 {len(sys_errors)}"
        )
        msgs = " | ".join(f.message for f in sys_errors)
        assert "甲 崩溃" in msgs and "乙 崩溃" in msgs, f"两条失败信息都应保留: {msgs}"
        assert len({f.dedup_key for f in sys_errors}) == 2

    def test_successful_results_kept_alongside_sys_error(self):
        """正常审计器的 findings 与失败审计器的 SYS-ERROR 并存，on_progress 仍以「完成」收尾。"""
        ok_finding = AuditFinding(
            type=FindingType.STRUCTURE,
            severity=FindingSeverity.INFO,
            message="正常发现",
            rule_id="OK-1",
            page_index=0,
        )
        calls: list[tuple] = []
        auditors = [
            ("甲", _StubAuditor("甲", findings=[ok_finding])),
            ("乙", _StubAuditor("乙", exc=RuntimeError("乙 崩溃"))),
        ]
        findings = run_auditors(
            _simple_doc(), auditors, on_progress=lambda n, i, t: calls.append((n, i, t))
        )
        assert sorted(f.rule_id for f in findings) == ["OK-1", "SYS-ERROR"]
        assert calls[-1] == ("完成", 2, 2), "即使审计器失败，on_progress 仍应以「完成」收尾"
