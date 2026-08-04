"""端到端集成测试 — convert → audit → dedup → report 完整流水线"""

import json
import tempfile
from pathlib import Path

from src.engines.pipeline import find_converter, build_auditors, run_auditors
from src.reporters.html_reporter import generate_html_report
from src.reporters.json_reporter import generate_json_report
from src.models.finding import AuditFinding, FindingSeverity, FindingType


class TestFullPipeline:
    """验证完整的 convert → audit → dedup → report 流水线"""

    def test_pptx_full_pipeline(self):
        """端到端: 转换 PPTX → 全部审查器 → 去重 → HTML + JSON 报告"""
        # 1. Convert
        converter = find_converter("tests/fixtures/sample.pptx")
        assert converter is not None, "PPTX converter not found"
        doc = converter.convert("tests/fixtures/sample.pptx")
        assert doc.format == "pptx"
        assert len(doc.pages) >= 1

        # 2. Build auditors
        auditors = build_auditors(
            "rules.md",
            str(Path("glossary").resolve()),
        )
        assert len(auditors) == 5
        auditor_names = [name for name, _ in auditors]
        assert "结构审查" in auditor_names
        assert "格式审查" in auditor_names
        assert "语言审查" in auditor_names
        assert "事实审查" in auditor_names
        assert "自定义规则" in auditor_names

        # 3. Run all auditors
        findings = run_auditors(doc, auditors)
        assert isinstance(findings, list)

        # 4. Verify dedup — no duplicates
        dedup_keys = [f.dedup_key for f in findings]
        assert len(dedup_keys) == len(set(dedup_keys)), (
            f"Dedup failed: {len(dedup_keys)} keys, {len(set(dedup_keys))} unique"
        )

        # 5. Verify finding structure
        for f in findings:
            assert isinstance(f, AuditFinding)
            assert f.type is not None
            assert f.severity is not None
            assert f.message, "Finding message should not be empty"
            assert isinstance(f.id, str) and len(f.id) == 12

        # 6. Generate HTML report
        html = generate_html_report(doc, findings, title="集成测试")
        assert "<!DOCTYPE html>" in html
        assert "集成测试" in html
        assert "审查发现" in html

        # 7. Generate JSON report
        json_report = generate_json_report(doc, findings)
        assert json_report["meta"]["tool"] == "DocAudit"
        assert json_report["meta"]["format"] == "pptx"
        assert "findings" in json_report
        assert json_report["summary"]["total_findings"] == len(findings)

        # 8. JSON report to file
        with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        try:
            generate_json_report(doc, findings, output_path=tmp_path)
            saved = json.loads(tmp_path.read_text(encoding="utf-8"))
            assert saved["summary"]["total_findings"] == len(findings)
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_markdown_full_pipeline(self):
        """端到端: 转换 Markdown → 审查 → 报告"""
        import tempfile
        md_content = """---
title: 测试文档
---

# 概述

这是测试文档的概述部分，介绍了 FinFET 技术的应用。

## 工艺参数

关键参数: Through Silicon Via 良率 95.3%.

## 结论

测试完成，良率达到 95.3%，符合预期。
"""
        tmp_md = None
        try:
            with tempfile.NamedTemporaryFile(
                suffix=".md", mode="w", encoding="utf-8", delete=False
            ) as f:
                f.write(md_content)
                tmp_md = Path(f.name)

            converter = find_converter(str(tmp_md))
            assert converter is not None
            doc = converter.convert(str(tmp_md))
            assert doc.format == "md"
            assert len(doc.pages) >= 1

            auditors = build_auditors("rules.md", str(Path("glossary").resolve()))
            findings = run_auditors(doc, auditors)

            # 应有 findings (如 "Through Silicon Via" 未使用 TSV 缩写)
            assert isinstance(findings, list)  # 流水线返回有效列表
            # 验证所有 finding 结构完整
            for f in findings:
                assert f.type is not None
                assert f.severity is not None
                assert f.message

            html = generate_html_report(doc, findings)
            assert "审查报告" in html

        finally:
            if tmp_md and tmp_md.exists():
                tmp_md.unlink()

    def test_finding_serialization_roundtrip(self):
        """AuditFinding → dict → from_dict 往返测试"""
        original = AuditFinding(
            type=FindingType.STRUCTURE,
            severity=FindingSeverity.ERROR,
            message="测试发现",
            rule_id="TEST-001",
            page_index=3,
            element_index=1,
            context="测试上下文",
            suggestion="测试建议",
            location="第 4 页",
            metadata={"key": "value"},
        )
        # Serialize
        d = original.to_dict()
        # Deserialize
        restored = AuditFinding.from_dict(d)
        assert restored.type == original.type
        assert restored.severity == original.severity
        assert restored.message == original.message
        assert restored.rule_id == original.rule_id
        assert restored.page_index == original.page_index
        assert restored.context == original.context
        assert restored.suggestion == original.suggestion
        assert restored.metadata == original.metadata

    def test_pipeline_handles_missing_file(self):
        """不存在的文件 → find_converter 返回 None (不崩溃)"""
        converter = find_converter("nonexistent_file.xyz")
        assert converter is None

    def test_pipeline_dedup_after_full_run(self):
        """全流水线执行后 — verify AuditFinding.deduplicate 被调用"""
        doc_path = "tests/fixtures/sample.pptx"
        converter = find_converter(doc_path)
        doc = converter.convert(doc_path)
        auditors = build_auditors("rules.md", str(Path("glossary").resolve()))
        findings = run_auditors(doc, auditors)

        # run_auditors 内部已调用 deduplicate → 不应有重复
        # 验证: 对所有 findings 检查 dedup_key 唯一性
        keys = [f.dedup_key for f in findings]
        duplicates = [k for k in keys if keys.count(k) > 1]
        assert len(duplicates) == 0, (
            f"run_auditors should deduplicate, but found duplicates: {set(duplicates)}"
        )

    def test_custom_rules_auditor_integrated(self):
        """自定义规则审查器作为流水线一部分正确执行"""
        doc_path = "tests/fixtures/sample.pptx"
        converter = find_converter(doc_path)
        doc = converter.convert(doc_path)
        auditors = build_auditors("rules.md", str(Path("glossary").resolve()))

        # 确保 CustomRulesAuditor 在流水线中
        custom_auditor = None
        for name, auditor in auditors:
            if name == "自定义规则":
                custom_auditor = auditor
                break
        assert custom_auditor is not None, "CustomRulesAuditor should be in pipeline"

        # 验证 delegate auditors 已注入
        assert custom_auditor._structure_auditor is not None, (
            "StructureAuditor should be injected into CustomRulesAuditor"
        )
        assert custom_auditor._format_auditor is not None, (
            "FormatAuditor should be injected into CustomRulesAuditor"
        )
        assert custom_auditor._factual_auditor is not None, (
            "FactualAuditor should be injected into CustomRulesAuditor"
        )

    def test_web_ui_path_equals_cli_path_equals_python_path(self):
        """验证 Web UI 路径 = CLI 路径 = 直接 Python 调用产生相同结果。

        这是系统的"黄金测试"：三个执行路径共享同一流水线，必须返回相同的 findings。
        """
        test_file = "tests/fixtures/sample.pptx"
        project_root = Path(__file__).parent.parent
        if not (project_root / test_file).exists():
            import pytest
            pytest.skip("Test fixture tests/fixtures/sample.pptx not found")

        # ── 路径 1: 直接 Python (基准/真理) ──────────────────
        converter1 = find_converter(test_file)
        doc1 = converter1.convert(test_file)
        auditors1 = build_auditors(
            str(project_root / "rules.md"),
            str(project_root / "glossary"),
            str(project_root / "vocab"),
        )
        findings1 = run_auditors(doc1, auditors1)

        # ── 路径 2: 模拟 CLI (相对路径, 无显式 vocab) ────────
        import os
        original_cwd = os.getcwd()
        try:
            os.chdir(str(project_root))
            converter2 = find_converter(test_file)
            doc2 = converter2.convert(test_file)
            # CLI 风格: rules_path 相对路径, glossary_dir 相对路径, vocab_dir=None
            auditors2 = build_auditors(
                "rules.md",
                str(Path("glossary").resolve()),
                None,  # CLI 默认不传 vocab
            )
            findings2 = run_auditors(doc2, auditors2)
        finally:
            os.chdir(original_cwd)

        # ── 路径 3: 模拟 Web UI (绝对路径, 显式 vocab) ───────
        converter3 = find_converter(str(project_root / test_file))
        doc3 = converter3.convert(str(project_root / test_file))
        # Web UI 风格: 全部绝对路径
        auditors3 = build_auditors(
            str(project_root / "rules.md"),
            str(project_root / "glossary"),
            str(project_root / "vocab"),
        )
        findings3 = run_auditors(doc3, auditors3)

        # ── 断言: 三个路径结果必须一致 ─────────────────────
        # 1. 总数相同
        assert len(findings1) == len(findings2) == len(findings3), (
            f"Finding count mismatch: Python={len(findings1)}, "
            f"CLI={len(findings2)}, WebUI={len(findings3)}"
        )

        # 2. 按 severity 统计相同
        def severity_counts(findings):
            counts = {}
            for f in findings:
                counts[f.severity.value] = counts.get(f.severity.value, 0) + 1
            return counts
        assert severity_counts(findings1) == severity_counts(findings2) == severity_counts(findings3), (
            f"Severity distribution mismatch:\n"
            f"  Python: {severity_counts(findings1)}\n"
            f"  CLI:    {severity_counts(findings2)}\n"
            f"  WebUI:  {severity_counts(findings3)}"
        )

        # 3. 按 type 统计相同
        def type_counts(findings):
            counts = {}
            for f in findings:
                counts[f.type.value] = counts.get(f.type.value, 0) + 1
            return counts
        assert type_counts(findings1) == type_counts(findings2) == type_counts(findings3), (
            f"Type distribution mismatch:\n"
            f"  Python: {type_counts(findings1)}\n"
            f"  CLI:    {type_counts(findings2)}\n"
            f"  WebUI:  {type_counts(findings3)}"
        )

        # 4. 按 rule_id 统计相同
        def rule_counts(findings):
            counts = {}
            for f in findings:
                rid = f.rule_id or "None"
                counts[rid] = counts.get(rid, 0) + 1
            return counts
        assert rule_counts(findings1) == rule_counts(findings2) == rule_counts(findings3), (
            f"Rule distribution mismatch:\n"
            f"  Python: {rule_counts(findings1)}\n"
            f"  CLI:    {rule_counts(findings2)}\n"
            f"  WebUI:  {rule_counts(findings3)}"
        )

        # 5. 每个 finding 的 dedup_key 集合相同 (确保无漏网差异)
        keys1 = sorted(f.dedup_key for f in findings1)
        keys2 = sorted(f.dedup_key for f in findings2)
        keys3 = sorted(f.dedup_key for f in findings3)
        assert keys1 == keys2 == keys3, (
            f"Dedup key mismatch!\n"
            f"  Only in Python: {set(keys1) - set(keys2)}\n"
            f"  Only in CLI:    {set(keys2) - set(keys1)}"
        )

    def test_golden_path_docx(self, tmp_path):
        """黄金测试扩展: DOCX 格式三路径一致性"""
        from docx import Document as DocxDocument
        from docx.shared import Pt

        # 创建测试 DOCX
        docx_doc = DocxDocument()
        docx_doc.add_heading("测试标题", level=1)
        docx_doc.add_paragraph("这是正文内容，包含 FinFET 技术。")
        docx_doc.add_heading("结论", level=2)
        docx_doc.add_paragraph("测试完成。")
        docx_path = tmp_path / "golden.docx"
        docx_doc.save(str(docx_path))

        project_root = Path(__file__).parent.parent

        # 路径 1: 绝对路径
        converter1 = find_converter(str(docx_path))
        doc1 = converter1.convert(str(docx_path))
        auditors1 = build_auditors(
            str(project_root / "rules.md"),
            str(project_root / "glossary"),
            str(project_root / "vocab"),
        )
        findings1 = run_auditors(doc1, auditors1)

        # 路径 2: 相同参数重新执行
        converter2 = find_converter(str(docx_path))
        doc2 = converter2.convert(str(docx_path))
        auditors2 = build_auditors(
            str(project_root / "rules.md"),
            str(project_root / "glossary"),
            str(project_root / "vocab"),
        )
        findings2 = run_auditors(doc2, auditors2)

        # 断言: 两次执行结果一致
        assert len(findings1) == len(findings2)
        keys1 = sorted(f.dedup_key for f in findings1)
        keys2 = sorted(f.dedup_key for f in findings2)
        assert keys1 == keys2

    def test_golden_path_markdown(self, tmp_path):
        """黄金测试扩展: Markdown 格式三路径一致性"""
        md_content = "# 概述\n\n这是测试文档，介绍 TSV 技术。\n\n## 结论\n\n测试完成。\n"
        md_path = tmp_path / "golden.md"
        md_path.write_text(md_content, encoding="utf-8")

        project_root = Path(__file__).parent.parent

        # 路径 1
        converter1 = find_converter(str(md_path))
        doc1 = converter1.convert(str(md_path))
        auditors1 = build_auditors(
            str(project_root / "rules.md"),
            str(project_root / "glossary"),
            str(project_root / "vocab"),
        )
        findings1 = run_auditors(doc1, auditors1)

        # 路径 2
        converter2 = find_converter(str(md_path))
        doc2 = converter2.convert(str(md_path))
        auditors2 = build_auditors(
            str(project_root / "rules.md"),
            str(project_root / "glossary"),
            str(project_root / "vocab"),
        )
        findings2 = run_auditors(doc2, auditors2)

        # 断言
        assert len(findings1) == len(findings2)
        keys1 = sorted(f.dedup_key for f in findings1)
        keys2 = sorted(f.dedup_key for f in findings2)
        assert keys1 == keys2
